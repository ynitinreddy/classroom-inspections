import base64
import datetime
import io
import os

import numpy as np
import streamlit as st
from PIL import Image
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from ultralytics import YOLO
import torch
import openai

# -----------------------------------------------------------
# 0️⃣  Environment & API Keys
# -----------------------------------------------------------
load_dotenv(dotenv_path=r"C:\Users\yniti\Downloads\classroom_inspector_api_key.env")
openai.api_key = os.getenv("OPENAI_API_KEY")

# -----------------------------------------------------------
# 1️⃣  Streamlit Page Config & Global Styling
# -----------------------------------------------------------
st.set_page_config(page_title="AI Classroom Inspector", layout="wide")

st.markdown(
    """
    <style>
    .reportview-container .main .block-container {
        padding: 2rem;
        background-color: #ffffff;
        color: #000000;
    }
    .stButton>button {
        color: white;
        background-color: #8C1D40;            /* ASU maroon */
        border: none;
        padding: 0.5rem 1rem;
        font-size: 16px;
        border-radius: 8px;
    }
    .stButton>button:hover {
        background-color: #FFC627;             /* ASU gold */
        color: black;
    }
    .stSelectbox label, .stCheckbox label, .stTextArea label,
    .stSubheader, .stCaption, .stMarkdown, .stTextInput label {
        color: #8C1D40 !important;
    }
    .stTextInput>div>input, .stTextArea textarea {
        background-color: #fff8dc;             /* very‑light gold */
        border: 1px solid #8C1D40;
        color: black;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# -----------------------------------------------------------
# 2️⃣  Session‑State Defaults (remember previous choices)
# -----------------------------------------------------------
DEFAULT_INSPECTOR = "Nitin"
DEFAULT_MODEL     = "Basic (fastest, cheapest)"
DEFAULT_YOLO      = True

if "inspector_name" not in st.session_state:
    st.session_state.inspector_name = DEFAULT_INSPECTOR
if "model_choice" not in st.session_state:
    st.session_state.model_choice = DEFAULT_MODEL
if "enable_yolo" not in st.session_state:
    st.session_state.enable_yolo = DEFAULT_YOLO

# -----------------------------------------------------------
# 3️⃣  Utility Functions
# -----------------------------------------------------------

def get_image_base64(path: str) -> str:
    """Return base64 string of a local image for inline HTML/Markdown."""
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()


def image_to_base64(image_file) -> str:
    img = Image.open(image_file).convert("RGB")
    buff = io.BytesIO()
    img.save(buff, format="JPEG")
    return base64.b64encode(buff.getvalue()).decode()


def load_yolo_cached():
    @st.cache_resource(show_spinner="Loading YOLO model …")
    def _load():
        return YOLO("best_yolov8n.pt")  # ~6 MB, CPU‑friendly

    return _load()


def call_gpt_hybrid(images, prompt, model, anomaly_data=None):
    """
    1) Build a complete YOLO summary (zeros for absent classes)
    2) Prepend it to your human prompt
    3) Send that single text block + images to GPT
    """
    # 1) Build a full counts table (use your CUSTOM_CLASSES mapping)
    labels = list(CUSTOM_CLASSES.values())
    if anomaly_data:
        full_counts = {label: anomaly_data.get(label, 0) for label in labels}
    else:
        full_counts = {label: 0 for label in labels}
    summary = "\n".join(f"- {label}: {count}" for label, count in full_counts.items())

    # 2) Inject the summary explicitly (and forbid guessing)
    injection = (
        "YOLO Detection Summary (use these counts to decide Present vs Absent – do NOT guess):\n"
        f"{summary}\n\n"
    )
    full_text = injection + prompt

    # 3) Build the multimodal blocks
    blocks = [{"type": "text", "text": full_text}]
    for img in images:
        blocks.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{image_to_base64(img)}"},
        })

    # 4) Call the OpenAI API
    client = openai.OpenAI()
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": blocks}],
        max_tokens=1000,
        temperature=0.1,
        top_p=0.9,
    )
    return resp.choices[0].message.content



def generate_docx_report(
    report_text,
    original_images,
    anomaly_images=None,
    class_number: str | None = None,
    inspector_name: str | None = None,
):
    """Return (BytesIO, filename, local_path) for a full DOCX report."""

    doc = Document()

    # Header info
    doc.add_heading("Classroom Inspection Report", 0)
    doc.add_paragraph(f"Class Number: {class_number if class_number else '__________________'}")
    doc.add_paragraph(f"Inspector: {inspector_name if inspector_name else '__________________'}")
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # Inspection summary
    doc.add_heading("1. Inspection Summary", level=1)
    for line in report_text.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip(), style="List Bullet")

    # Original images
    doc.add_heading("2. Uploaded Classroom Images", level=1)
    for i, img_file in enumerate(original_images):
        img = Image.open(img_file)
        img_io = io.BytesIO()
        img.save(img_io, format="JPEG")
        img_io.seek(0)
        doc.add_paragraph(f"Original Image {i + 1}")
        doc.add_picture(img_io, width=Inches(5))
        doc.add_paragraph("")

    # Anomaly detections
    if anomaly_images:
        doc.add_heading("3. YOLO Enhanced Detections", level=1)
        for i, img in enumerate(anomaly_images):
            img_io = io.BytesIO()
            img.save(img_io, format="JPEG")
            img_io.seek(0)
            doc.add_paragraph(f"Recognized Image {i + 1}")
            doc.add_picture(img_io, width=Inches(5))
            doc.add_paragraph("")

    # Build filename
    today_str = datetime.date.today().strftime("%Y-%m-%d")
    inspector_part = (
        inspector_name.strip().replace(" ", "_") if inspector_name else "anonymous"
    )
    classroom_part = (
        class_number.strip().replace(" ", "_") if class_number else "unknownclass"
    )
    file_name = f"{today_str}_{inspector_part}_{classroom_part}_report.docx"

    # Save to BytesIO and local temp dir
    output_io = io.BytesIO()
    doc.save(output_io)
    output_io.seek(0)

    os.makedirs("temp_reports", exist_ok=True)
    local_path = os.path.join("temp_reports", file_name)
    with open(local_path, "wb") as fp:
        doc.save(fp)

    return output_io, file_name, local_path


# -----------------------------------------------------------
# 4️⃣  Header & Logo
# -----------------------------------------------------------
logo_b64 = get_image_base64("ASU-logo.png")
st.markdown(
    f"<img src='data:image/png;base64,{logo_b64}' width='250'/>",
    unsafe_allow_html=True,
)
st.title("AI‑Powered Classroom Inspection · ASU Edition")
st.markdown("Upload classroom images to automatically detect issues and generate an inspection report.")

# -----------------------------------------------------------
# 5️⃣  Image Uploader
# -----------------------------------------------------------

st.subheader("Step 1: Upload Classroom Images")
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = "uploader_0"

uploaded_files = st.file_uploader(
    "Upload classroom images (multiple angles preferred)",
    accept_multiple_files=True,
    type=["jpg", "jpeg", "png"],
    key=st.session_state.uploader_key,
)

if uploaded_files:
    st.session_state.uploaded_files = uploaded_files

if st.session_state.uploaded_files:
    if st.button("🗑️ Clear Uploaded Images"):
        st.session_state.uploaded_files = []
        new_id = int(st.session_state.uploader_key.split("_")[1]) + 1
        st.session_state.uploader_key = f"uploader_{new_id}"
        st.rerun()

# -----------------------------------------------------------
# 6️⃣  Class No. & Inspector (Combined Step 1.5)
# -----------------------------------------------------------

st.subheader("Step 1.5: Inspection Details")

col_a, col_b = st.columns(2)
with col_a:
    class_number = st.text_input(
        "Classroom Number (e.g., 'DH 101') – leave blank if unknown",
        value=st.session_state.get("class_number", ""),
        key="class_number",
    )

with col_b:
    inspector_options = ["Nitin", "Jose", "Priyam", "Tanvi", "Others"]
    inspector_name = st.selectbox(
        "Inspector Name",
        inspector_options,
        index=inspector_options.index(st.session_state.inspector_name)
        if st.session_state.inspector_name in inspector_options
        else 0,
        key="inspector_name",
    )

# If "Others", prompt for custom name
custom_name = ""
if inspector_name == "Others":
    custom_name = st.text_input(
        "Enter inspector name:",
        value=st.session_state.get("custom_inspector_name", ""),
        key="custom_inspector_name",
    )

inspector_used = (
    custom_name.strip() if inspector_name == "Others" else inspector_name
)
if inspector_used == "":
    inspector_used = "anonymous"

# -----------------------------------------------------------
# 7️⃣  Model & YOLO Options (remembered)
# -----------------------------------------------------------

st.subheader("Step 2: Choose Model & Options")
model_options = [
    "Best (faster, lower cost)",
    "Basic (fastest, cheapest)",
    "Expert (most advanced reasoning for images) – need to add",
]
model_choice = st.selectbox(
    "LLM Model:",
    model_options,
    index=model_options.index(st.session_state.model_choice)
    if st.session_state.model_choice in model_options
    else model_options.index(DEFAULT_MODEL),
    key="model_choice",
)

enable_yolo_checkbox = st.checkbox(
    "Detect and highlight unusual objects?",
    value=st.session_state.enable_yolo,
    key="enable_yolo",
)

# After widgets, capture the current selections without touching their keys
enable_yolo = st.session_state.enable_yolo  # current checkbox state

# -----------------------------------------------------------
# 8️⃣  Model Mapping & YOLO Init (lazy)
# -----------------------------------------------------------

model_map = {
    "Best (faster, lower cost)": ("gpt-4o", "Using best model."),
    "Basic (fastest, cheapest)": ("gpt-4o-mini", "Using smaller model."),
    "Expert (most advanced reasoning for images) – need to add": (
        "gpt-4o",
        "Using expert-level reasoning model (gpt-4o).",
    ),
}
selected_model, model_comment = model_map[model_choice]

if st.session_state.enable_yolo:
    yolo_model = load_yolo_cached()
    # ONLY use this if those indexes are *really* what your .pt was trained on:
    # CUSTOM_CLASSES = {
    #     0: "911 Address",
    #     1: "Bill of Rights Constitution",
    #     2: "Capacity Sign",
    #     3: "Classroom Layout",
    #     4: "Classroom Support Pocket",
    #     5: "Clock",
    #     6: "Dirty Whiteboard",
    #     7: "ERG Layout",
    #     8: "Exit Sign",
    #     9: "Flag",
    #     10: "Miscellaneous Objects",
    #     11: "No Food/Drinks Sign",
    #     12: "Recycle Bin",
    #     13: "Scrapes",
    #     14: "Stains",
    #     15: "Trash Bin",
    #     16: "University Classrooms Pocket",
    #     17: "Whiteboard",
    #     18: "Window Covering",
    # }

    CUSTOM_CLASSES = {
        0: "911 Address",
        1: "Bill of Rights & Constitution",
        2: "Bins",
        3: "Capacity Sign",
        4: "Classroom Support Pocket",
        5: "Clock",
        6: "ERG",
        7: "Exit Sign",
        8: "Flag",
        9: "No Food or Drinks Sign",
        10: "Scuffs/Scrapes",
        11: "Stains",
        12: "UCL Pocket",
        13: "Whiteboard",
    }

    def detect_objects(images):
        counts: dict[str,int] = {}
        annotated: list[Image.Image] = []
        for img_file in images:
            img = Image.open(img_file).convert("RGB")
            # filter to exactly the IDs in your map
            results = yolo_model(np.array(img), classes=list(CUSTOM_CLASSES.keys()), conf=0.25, iou=0.30)


            for result in results:
                if not result.boxes:
                    continue
                # result.names is the list of labels from your .pt
                for box in result.boxes:
                    cls = int(box.cls[0])
                    label = CUSTOM_CLASSES[cls]
                    counts[label] = counts.get(label, 0) + 1

                # re‑plot with your model’s labels
                annotated.append(Image.fromarray(result.plot(conf=True, labels=True)))
        return counts, annotated

else:
    detect_objects = None



# -----------------------------------------------------------
# 9️⃣  Prompt Engineering
# -----------------------------------------------------------

def build_default_prompt(use_yolo: bool) -> str:
    extra = " and object counts" if use_yolo else ""
    return f"""
{model_comment}

You are a classroom inspection assistant. You will be given a set of classroom images and a list of objects detected by another vision model (YOLO){extra}. Use the following rules carefully:

🔒 VERY IMPORTANT:
- DO NOT guess object presence unless instructed below. If it’s not in the detection list and you aren’t explicitly told to "Visually inspect" it, say “Absent” or “Cannot determine.”
- Only describe condition details (e.g. cleanliness, damage, functionality) for items that you’re told to visually inspect.
- Be brief and specific.



1. **Side Walls** – Visually inspect.  
2. **Ceiling** – Visually inspect.  
3. **White Board** – Use YOLO to decide Present/Absent, then visually check for clean versus writing.  
4. **Floor** – Visually inspect.  
5. **Bins** – Use YOLO to decide Present/Absent, then visually count trash bins vs. recycle bins.  
6. **Exit Sign** – Rely **only** on YOLO for Present/Absent.  
7. **Lights** – Visually inspect.  
8. **Flag** – Rely **only** on YOLO for Present/Absent.  
9. **“No Food/Drinks” Plaque** – Rely **only** on YOLO for Present/Absent.  
10. **Instructor’s Desk** – Visually inspect for Present/Absent and note its condition (e.g., organized, cluttered).  
11. **Clock** – Rely **only** on YOLO for Present/Absent.  
12. **Capacity Sign** – Rely **only** on YOLO for Present/Absent.  
13. **UCL Pocket** – Rely **only** on YOLO for Present/Absent.  
14. **Classroom Support Pocket** – Rely **only** on YOLO for Present/Absent.  
15. **911 Address on Door Frame** – Rely **only** on YOLO for Present/Absent.  
16. **ERG** – Rely **only** on YOLO for Present/Absent.  
17. **Bill of Rights & Constitution** – Rely **only** on YOLO for Present/Absent.  
18. **Additional Comments** – Visually note anything odd (messy room, safety issues, etc.).

Be accurate. Follow these rules exactly.  
"""

prompt_default = build_default_prompt(st.session_state.enable_yolo)
with st.expander("⚙️ More Options: Edit Inspection Prompt"):
    prompt = st.text_area("LLM Prompt", prompt_default, height=260)

# -----------------------------------------------------------
# 🔟  Run Inspection
# -----------------------------------------------------------

center = st.columns([1, 2, 1])[1]
run_btn = center.button("🔍 Run Inspection", use_container_width=True)
status = st.empty()

if run_btn:
    if not st.session_state.uploaded_files:
        st.error("Please upload at least one image.")
        st.stop()

    # 1) Show uploaded images
    status.info("Preparing images…")
    with st.expander("📷 View uploaded images"):
        cols = st.columns(min(len(st.session_state.uploaded_files), 4))
        for i, f in enumerate(st.session_state.uploaded_files):
            with cols[i % 4]:
                st.image(Image.open(f), caption=f"Image {i + 1}", use_container_width=True)
    
    # st.write("📝 Model reports these classes:", load_yolo_cached().model.names)


    # 2) YOLO object detection (optional)
    detections = {}
    annotated_imgs = []
    if st.session_state.enable_yolo and detect_objects:
        status.info("Detecting objects with YOLO…")
        detections, annotated_imgs = detect_objects(st.session_state.uploaded_files)
        st.write("**Detected objects:**", detections or "No objects detected.")

        if annotated_imgs:
            with st.expander("📦 YOLO Detections"):
                cols = st.columns(min(len(annotated_imgs), 4))
                for i, im in enumerate(annotated_imgs):
                    with cols[i % 4]:
                        st.image(im, caption=f"Detection {i + 1}", use_container_width=True)

        # 3) Run GPT inspection over images + detection summary
        status.info("Analyzing classroom with AI Vision…")
        report = call_gpt_hybrid(
            st.session_state.uploaded_files,
            prompt,
            selected_model,
            anomaly_data=detections if st.session_state.enable_yolo else None,
        )


    # 4) Display the 13‑point inspection report
    st.subheader("📝 Inspection Report")
    st.markdown(report)

    # 5) Generate and offer download of DOCX report
    bio, file_name, local_file_path = generate_docx_report(
        report,
        st.session_state.uploaded_files,
        anomaly_images=annotated_imgs,
        class_number=class_number,
        inspector_name=inspector_used,
    )
    st.download_button("📄 Download DOCX Report", data=bio, file_name=file_name)

    status.success("All done! 🎉")


# -----------------------------------------------------------
# 1️⃣1️⃣  Sidebar – About
# -----------------------------------------------------------

with st.sidebar.expander("📄 About This Project"):
    avatar_b64 = get_image_base64("musk-photo-1.jpg")
    st.markdown(
        f"""
        <div style='text-align:center;'>
            <img src='data:image/jpeg;base64,{avatar_b64}' style='width:150px;border-radius:50%;'/>
            <div style='color:#8C1D40;font-size:16px;margin-top:8px;'><strong>Nitin Reddy Yarava</strong></div>
            <p style='font-size:16px;'>
                This project automates ASU classroom inspections using YOLOv8 for object detection and GPT‑4 Vision for reasoning.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

# -----------------------------------------------------------
# 1️⃣2️⃣  Footer
# -----------------------------------------------------------

st.markdown("---")
st.caption("Built by Nitin, a CS student at ASU ✌️")
